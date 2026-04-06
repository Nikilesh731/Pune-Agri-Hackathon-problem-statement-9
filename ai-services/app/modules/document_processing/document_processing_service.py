"""
Document Processing Service
Purpose: Clean orchestration layer over rebuilt schema, classifier, extraction service, and processor
"""
from typing import Optional, Dict, Any, List
import uuid
import sys
import os
import logging

# Runtime health check - CRITICAL for production stability
from app.modules.document_processing.runtime_health import (
    ensure_runtime_ready,
    get_available_tesseract_languages,
    get_runtime_health,
    get_tesseract_config,
    resolve_tesseract_binary,
)

from app.modules.document_processing.service_schema import DocumentProcessingRequest, DocumentProcessingResult
from app.modules.document_processing.classification_service import DocumentClassificationService
from app.modules.document_processing.extraction_service import DocumentExtractionService
from app.modules.document_processing.processors import DocumentProcessor
from app.modules.ai_assist.llm_assist_service import analyze_application
from app.modules.ai_assist.llm_refinement_service import add_llm_refinement_to_response
from app.modules.intelligence.intelligence_service import IntelligenceService
from app.modules.document_processing.ml_priority import predict_application_priority
from app.modules.document_processing.workflow_service import WorkflowService
from app.modules.document_processing.granite_extraction_service import GraniteExtractionService
from app.ml.ml_service import analyze_document_risk

ML_SERVICE_AVAILABLE = True


class DocumentProcessingService:
    """Clean orchestration service for document processing operations"""
    
    def __init__(self):
        # CRITICAL: Ensure runtime is ready before initializing services
        try:
            logger = logging.getLogger(__name__)
            get_runtime_health()
        except Exception as e:
            logger = logging.getLogger(__name__)
            logger.error(f"[HEALTH] Runtime not ready: {e}")
            raise RuntimeError(f"Document processing service cannot start: {e}")
        
        # Initialize rebuilt services
        self.classification_service = DocumentClassificationService()
        self.extraction_service = DocumentExtractionService()
        self.intelligence_service = IntelligenceService()
        self.workflow_service = WorkflowService()
        self.processor = DocumentProcessor(
            self.classification_service,
            self.extraction_service
        )
        
        # Lazy OCR service initialization - only when needed
        self._ocr_service = None
        self.paddle_ocr_service = None
        self.logger = logging.getLogger(__name__)

        # Initialize Granite extraction service with graceful degradation
        try:
            granite_endpoint = os.getenv("GRANITE_ENDPOINT")
            self.use_granite = os.getenv("USE_GRANITE", "true").lower() == "true"
            
            if granite_endpoint and self.use_granite:
                self.granite_service = GraniteExtractionService(granite_endpoint)
                logger.info("[INIT] Granite extraction service initialized successfully")
            else:
                if not granite_endpoint:
                    logger.info("[INIT] Granite: disabled (no endpoint configured)")
                if not self.use_granite:
                    logger.info("[INIT] USE_GRANITE=false - Granite service disabled")
                self.granite_service = None
        except Exception as e:
            logger.error(f"[INIT] Granite service initialization failed: {e}")
            logger.info("[INIT] Will fall back to traditional extraction logic")
            self.granite_service = None

    def _get_ocr_service(self):
        """Lazy initialization of OCR service"""
        self._ocr_service = None
        return None

    # --- Defensive helpers -------------------------------------------------
    def _safe_dict(self, val: Any) -> Dict[str, Any]:
        return val if isinstance(val, dict) else {}

    def _safe_list(self, val: Any) -> List[Any]:
        return val if isinstance(val, list) else []

    def _safe_str(self, val: Any) -> str:
        return val if isinstance(val, str) else ""

    def _safe_number(self, val: Any, default: float = 0.0) -> float:
        try:
            if isinstance(val, (int, float)):
                return float(val)
            if isinstance(val, str):
                s = val.strip().rstrip('%')
                import re
                s = re.sub(r"[^0-9.\.-]", "", s)
                if s:
                    return float(s)
        except Exception:
            pass
        return float(default)

    def _quality_issues(self, raw_text: str) -> List[str]:
        import re

        issues: List[str] = []
        text = re.sub(r"\s+", " ", raw_text or "").strip()

        if not text:
            issues.append("empty_text")
            return issues

        if len(text) < 50:
            issues.append("very_short_text")

        alpha_count = sum(1 for char in text if char.isalpha())
        symbol_count = sum(1 for char in text if not char.isalnum() and not char.isspace())
        if alpha_count < 5:
            issues.append("too_little_alphabetic_content")
        if symbol_count > max(10, int(len(text) * 0.4)):
            issues.append("mostly_symbols")

        return issues

    def _is_text_quality_poor(self, raw_text: str) -> bool:
        import re

        text = re.sub(r"\s+", " ", raw_text or "").strip()
        if not text or len(text) < 50:
            return True
        issues = self._quality_issues(raw_text)
        return any(issue in {"too_little_alphabetic_content", "mostly_symbols"} for issue in issues)

    def _merge_text_content(self, primary_text: str, secondary_text: str) -> str:
        import re

        primary_lines = [line.strip() for line in (primary_text or "").splitlines() if line.strip()]
        secondary_lines = [line.strip() for line in (secondary_text or "").splitlines() if line.strip()]
        merged_lines: List[str] = []
        seen_lines = set()

        for line in primary_lines + secondary_lines:
            normalized = re.sub(r"\s+", " ", line).strip()
            if not normalized or normalized in seen_lines:
                continue
            seen_lines.add(normalized)
            merged_lines.append(line)

        if not merged_lines:
            return (primary_text or secondary_text or "").strip()

        return "\n".join(merged_lines).strip()

    def _merge_docling_and_ocr_output(self, docling_output: Dict[str, Any], ocr_text: str) -> Dict[str, Any]:
        merged_output = dict(docling_output)
        merged_output["raw_text"] = self._merge_text_content(docling_output.get("raw_text", ""), ocr_text)

        metadata = self._safe_dict(merged_output.get("metadata", {}))
        metadata["ocr_fallback_used"] = True
        metadata["ocr_text_length"] = len(ocr_text)
        warnings = self._safe_list(metadata.get("warnings", []))
        if "ocr_fallback_used" not in warnings:
            warnings.append("ocr_fallback_used")
        metadata["warnings"] = warnings
        merged_output["metadata"] = metadata

        return merged_output

    def _resolve_tesseract_binary(self):
        """Resolve a usable Tesseract binary path for OCR fallback."""
        return resolve_tesseract_binary()

    def _get_ocr_language_candidates(self) -> List[str]:
        available_languages = set(get_available_tesseract_languages())
        preferred_languages = [language for language in ("eng", "hin", "mar") if language in available_languages]

        if preferred_languages:
            candidates = ["+".join(preferred_languages)]
        else:
            candidates = ["eng"]

        if "eng" not in preferred_languages:
            candidates.append("eng")

        return list(dict.fromkeys(candidates))

    def _build_canonical_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Build canonical data from structured_data with normalization
        
        Args:
            data: Dictionary containing structured_data
            
        Returns:
            Normalized canonical dictionary with non-empty values only
        """
        canonical = {}
        structured_data = self._safe_dict(data.get("structured_data", {}))
        
        for key, value in structured_data.items():
            if value is None:
                continue
            
            if isinstance(value, str):
                value = value.strip()
                if not value:
                    continue
            
            # Normalize amount-like fields to numeric
            if key in ["requested_amount", "claim_amount", "subsidy_amount", "amount"]:
                try:
                    normalized_amount = self._safe_number(value, default=0.0)
                    if normalized_amount > 0:
                        canonical[key] = normalized_amount
                    continue
                except Exception:
                    pass
            
            canonical[key] = value
        
        return canonical
    
    def _normalize_extracted_fields(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Normalize extracted_fields to ensure it's not empty when structured_data has values
        
        Args:
            data: Dictionary containing structured_data and extracted_fields
            
        Returns:
            Normalized extracted_fields dictionary
        """
        extracted_fields = self._safe_dict(data.get("extracted_fields", {}))
        structured_data = self._safe_dict(data.get("structured_data", {}))
        
        # Seed extracted_fields from structured_data if empty
        for key, value in structured_data.items():
            if key not in extracted_fields and value is not None and str(value).strip():
                extracted_fields[key] = value
        
        return extracted_fields

    def _build_officer_summary(self, extracted_data: Dict[str, Any]) -> str:
        """Build a deterministic, officer-facing summary from extracted fields."""
        structured = self._safe_dict(extracted_data.get("structured_data", {}))
        name = structured.get("farmer_name") or structured.get("applicant_name") or structured.get("applicant") or "Unknown"
        doc_type = extracted_data.get("document_type", "document").replace("_", " ")
        location = structured.get("location") or structured.get("village") or ""
        # requested amount if present
        amount = None
        for f in ["requested_amount", "claim_amount", "subsidy_amount", "amount"]:
            if f in structured and structured.get(f):
                amount = structured.get(f)
                break

        parts = []
        parts.append(f"{name} submitted a {doc_type}.")
        if location:
            parts.append(f"Location: {location}.")
        if amount:
            try:
                amt = self._safe_number(amount)
                parts.append(f"Requested amount extracted: {int(amt) if amt>=1 else amt}.")
            except Exception:
                parts.append(f"Requested amount extracted: {amount}.")

        # If scheme/subsidy specifics are missing, call that out
        if doc_type and ("subsid" in doc_type or "scheme" in doc_type):
            parts.append("Key identity fields are available; subsidy type or application details may be missing.")

        return " ".join(parts)
    
    def _extract_with_docling_granite(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract document information using Docling + Granite pipeline
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            
        Returns:
            Schema-compliant extraction result
            
        Raises:
            ValueError: If extraction fails
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Check if services are available
        if not self.docling_service:
            raise ValueError("Docling service not available")
        
        if not self.granite_service:
            raise ValueError("Granite service not available")
        
        try:
            # Step 1: Docling ingestion
            logger.info(f"[DOCLING] Starting ingestion for {filename}")
            docling_output = self.docling_service.ingest_document(file_data, filename)
            
            # Validate Docling output
            if not docling_output:
                raise ValueError("Docling ingestion returned no result")
            
            docling_text = docling_output.get('raw_text', '')
            logger.info(f"[DOCLING] success for {filename}: {len(docling_text)} chars")

            if self._is_text_quality_poor(docling_text):
                logger.warning(f"[OCR] fallback used for {filename}")
                try:
                    ocr_text = self._extract_text_from_file(file_data, filename)
                    if not ocr_text.strip() or self._is_text_quality_poor(ocr_text):
                        raise ValueError("OCR fallback produced poor text")
                    docling_output = self._merge_docling_and_ocr_output(docling_output, ocr_text)
                    logger.info(f"[PIPELINE] final text length: {len(docling_output.get('raw_text', ''))}")
                except Exception as ocr_error:
                    logger.error(f"[OCR] failed for {filename}: {ocr_error}")
                    failure_payload = self._create_processing_failure_payload(
                        filename,
                        str(ocr_error),
                        "OCR_FAILURE"
                    )
                    return failure_payload
            
            # Step 2: Granite extraction
            logger.info(f"[GRANITE] Starting extraction for {filename}")
            granite_output = self.granite_service.extract_with_granite(docling_output)
            
            # Validate Granite output
            if not granite_output:
                raise ValueError("Granite extraction returned no results")
            
            logger.info(f"[GRANITE] Extraction complete: {granite_output.get('document_type', 'unknown')}")
            
            # Add processing metadata
            granite_output['processing_metadata'] = {
                'pipeline': 'docling_granite',
                'docling_available': True,
                'granite_available': True,
                'file_type': docling_output.get('metadata', {}).get('file_type', 'unknown'),
                'docling_text_length': len(docling_output.get('raw_text', '')),
                'has_tables': len(docling_output.get('tables', [])) > 0,
                'has_sections': len(docling_output.get('sections', [])) > 0
            }

            logger.info(f"[PIPELINE] final text length: {len(docling_output.get('raw_text', ''))}")
            
            return granite_output
            
        except Exception as e:
            logger.error(f"[DOCLING-GRANITE] Pipeline failed for {filename}: {e}")
            raise ValueError(f"Docling+Granite extraction failed: {e}")
    
    def _should_use_docling_granite(self, filename: str) -> bool:
        """
        Determine if enhanced pipeline should be used for this file
        
        Args:
            filename: Original filename
            
        Returns:
            False - using PyMuPDF/OCR pipeline instead
        """
        # Disabled - using PyMuPDF/OCR pipeline exclusively
        return False
    
    def _create_processing_failure_payload(self, filename: str, error_message: str, failure_type: str = "PROCESSING_FAILURE") -> Dict[str, Any]:
        """
        Create a safe failure payload for processing errors
        
        Args:
            filename: Original filename
            error_message: Error message
            failure_type: Type of failure (OCR_FAILURE, PROCESSING_FAILURE)
            
        Returns:
            Schema-compliant failure payload
        """
        return {
            "document_type": "unknown",
            "structured_data": {},
            "extracted_fields": {},
            "missing_fields": [],
            "confidence": 0.0,
            "reasoning": [f"{failure_type}: {error_message}"],
            "classification_confidence": 0.0,
            "classification_reasoning": {
                "keywords_found": [],
                "structural_indicators": [],
                "confidence_factors": [f"{failure_type}: {error_message}"]
            },
            "risk_flags": [
                {
                    "code": failure_type,
                    "severity": "high",
                    "message": f"{failure_type}: {error_message}"
                }
            ],
            "decision_support": {
                "decision": "manual_review_required",
                "confidence": 0.0,
                "reasoning": [f"{failure_type}: {error_message}"]
            },
            "canonical": {},
            "summary": f"Document processing failed: {failure_type} - {error_message}",
            "ai_summary": f"Document processing failed: {failure_type} - {error_message}",
            "processing_metadata": {
                "pipeline": "docling_granite",
                "processing_failure": True,
                "failure_type": failure_type,
                "error_message": error_message
            }
        }
    
    async def process_document(
        self,
        file_data: bytes,
        filename: str,
        processing_type: str = "full_process",
        options: Optional[Dict[str, Any]] = None
    ) -> DocumentProcessingResult:
        """
        Process document based on specified type
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            processing_type: Type of processing to apply
            options: Additional processing options
        
        Returns:
            DocumentProcessingResult with processing results
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Ensure options exists
        options = options or {}
        
        # PART 1: REQUIRED TEXT EXTRACTION ENTRY STEP
        # If valid OCR text already exists, use it; otherwise extract from file
        existing_ocr_text = options.get("ocr_text")
        if existing_ocr_text and isinstance(existing_ocr_text, str) and existing_ocr_text.strip():
            logger.info(f"[EXTRACT] Using existing OCR text: {len(existing_ocr_text)} chars")
        else:
            # Extract text from binary file before workflow
            if file_data:
                try:
                    # NEW: Check if we should use Docling+Granite pipeline
                    if self._should_use_docling_granite(filename):
                        logger.info(f"[EXTRACT] Using Docling+Granite pipeline for {filename}")
                        try:
                            # Use Docling+Granite for full extraction
                            granite_result = self._extract_with_docling_granite(file_data, filename)
                            
                            # Convert to result format and return early
                            # This bypasses the traditional processor since Granite provides full schema
                            return DocumentProcessingResult(
                                request_id=str(uuid.uuid4()),
                                success=True,
                                processing_time_ms=0,  # TODO: Add timing
                                processing_type=processing_type,
                                filename=filename,
                                data=granite_result,
                                metadata={"pipeline": "docling_granite"},
                                error_message=None
                            )
                            
                        except Exception as docling_granite_error:
                            logger.warning(f"[EXTRACT] Docling+Granite pipeline failed for {filename}: {docling_granite_error}")
                            logger.warning(f"[EXTRACT] Falling back to traditional extraction")
                            
                            # TRUE FALLBACK: Use traditional text extraction
                            try:
                                extracted_text = self._extract_text_from_file(file_data, filename)
                                logger.info(f"[EXTRACT] Fallback extracted text length: {len(extracted_text)}")
                                
                                # PART 5: FAIL EXPLICITLY ON EMPTY EXTRACTION
                                if not extracted_text.strip():
                                    logger.error(f"[EXTRACT] Fallback empty text extraction for {filename}")
                                    return DocumentProcessingResult(
                                        request_id=str(uuid.uuid4()),
                                        success=False,
                                        processing_time_ms=0,
                                        processing_type=processing_type,
                                        filename=filename,
                                        data=None,
                                        metadata={},
                                        error_message="Unable to extract text from document - file may be corrupted or unreadable"
                                    )
                                
                                # PART 2: PASS EXTRACTED TEXT INTO OPTIONS
                                options["ocr_text"] = extracted_text
                                logger.info(f"[EXTRACT] Passing fallback extracted text into workflow")
                                
                            except Exception as fallback_error:
                                logger.error(f"[EXTRACT] Fallback also failed: {fallback_error}")
                                
                                # Create processing failure payload for complete failure
                                failure_payload = self._create_processing_failure_payload(
                                    filename, str(fallback_error), "COMPLETE_PROCESSING_FAILURE"
                                )
                                
                                return DocumentProcessingResult(
                                    request_id=str(uuid.uuid4()),
                                    success=True,  # Success=true but with failure flags for manual review
                                    processing_time_ms=0,
                                    processing_type=processing_type,
                                    filename=filename,
                                    data=failure_payload,
                                    metadata={"pipeline": "docling_granite", "processing_failure": True},
                                    error_message=None
                                )
                    
                    # FALLBACK: Use traditional text extraction
                    extracted_text = self._extract_text_from_file(file_data, filename)
                    logger.info(f"[EXTRACT] Extracted text length: {len(extracted_text)}")
                    
                    # PART 5: FAIL EXPLICITLY ON EMPTY EXTRACTION
                    if not extracted_text.strip():
                        logger.error(f"[EXTRACT] Empty text extraction for {filename}")
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=False,
                            processing_time_ms=0,
                            processing_type=processing_type,
                            filename=filename,
                            data=None,
                            metadata={},
                            error_message="Unable to extract text from document - file may be corrupted or unreadable"
                        )
                    
                    # PART 2: PASS EXTRACTED TEXT INTO OPTIONS
                    options["ocr_text"] = extracted_text
                    logger.info(f"[EXTRACT] Passing extracted text into workflow")
                    
                except Exception as extraction_error:
                    logger.error(f"[EXTRACT] Text extraction failed: {extraction_error}")
                    
                    # CHECK IF THIS IS AN OCR FAILURE (IMAGE FILE)
                    file_extension = filename.lower().split('.')[-1] if '.' in filename.lower() else ''
                    is_image_file = file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']
                    error_text = str(extraction_error)
                    is_ocr_failure = (
                        "OCR failed" in error_text
                        or "Tesseract" in error_text
                        or "OCR_FAILURE" in error_text
                    )
                    
                    if is_image_file and is_ocr_failure:
                        # OCR FAILURE: Return proper response with risk_flags
                        logger.error(f"[OCR] failed for image {filename}: {extraction_error}")
                        
                        # Create OCR failure response data
                        ocr_failure_data = {
                            "document_type": "unknown",
                            "structured_data": {},
                            "extracted_fields": {},
                            "missing_fields": [],
                            "confidence": 0.0,
                            "reasoning": ["OCR failed - image OCR unavailable or extraction failed"],
                            "classification_confidence": 0.0,
                            "classification_reasoning": {
                                "keywords_found": [],
                                "structural_indicators": [],
                                "confidence_factors": ["OCR failed - image OCR unavailable or extraction failed"]
                            },
                            "risk_flags": [
                                {
                                    "code": "OCR_FAILURE",
                                    "severity": "high",
                                    "message": "OCR failed - image OCR unavailable or extraction failed"
                                }
                            ],
                            "decision_support": {
                                "decision": "manual_review_required",
                                "confidence": 0.0,
                                "reasoning": ["OCR failed - image OCR unavailable or extraction failed"]
                            },
                            "canonical": {},
                            "summary": "Automatic extraction failed. Manual review required.",
                            "ai_summary": "Automatic extraction failed. Manual review required."
                        }
                        
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=True,  # Success=true but with OCR failure flags
                            processing_time_ms=0,
                            processing_type=processing_type,
                            filename=filename,
                            data=ocr_failure_data,
                            metadata={"ocr_failure": True},
                            error_message=None
                        )
                    else:
                        # REGULAR EXTRACTION FAILURE: Return normal error
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=False,
                            processing_time_ms=0,
                            processing_type=processing_type,
                            filename=filename,
                            data=None,
                            metadata={},
                            error_message=f"Text extraction failed: {str(extraction_error)}"
                        )
            else:
                logger.warning("[EXTRACT] No file data provided for extraction")
                return DocumentProcessingResult(
                    request_id=str(uuid.uuid4()),
                    success=False,
                    processing_time_ms=0,
                    processing_type=processing_type,
                    filename=filename,
                    data=None,
                    metadata={},
                    error_message="No file data provided for text extraction"
                )
        
        # PART 3: CALL PROCESSOR WITH EXTRACTED TEXT
        processing_result = self.processor.process_document_workflow(
            file_data, filename, processing_type, options
        )
        
        return self._convert_to_result_format(processing_result)
    
    async def process_batch(
        self,
        documents: List[Dict[str, Any]],
        processing_type: str = "full_process"
    ) -> List[DocumentProcessingResult]:
        """
        Process multiple documents in batch
        
        Args:
            documents: List of document data with metadata
            processing_type: Type of processing to apply
        
        Returns:
            List of DocumentProcessingResult objects
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Set processing type for each document unless explicitly provided
        for doc in documents:
            if 'processing_type' not in doc:
                doc['processing_type'] = processing_type
        
        # PART 6: ENSURE ALL PROCESSING PATHS USE THE SAME FIX
        # Apply text extraction to each document before batch processing
        processed_documents = []
        for doc in documents:
            file_data = doc.get('file_data')
            filename = doc.get('filename', f'document_{len(processed_documents)+1}')
            options = doc.get('options', {})
            
            # Ensure options exists
            if not options:
                options = {}
                doc['options'] = options
            
            # PART 1: REQUIRED TEXT EXTRACTION ENTRY STEP
            existing_ocr_text = options.get("ocr_text")
            if existing_ocr_text and isinstance(existing_ocr_text, str) and existing_ocr_text.strip():
                logger.info(f"[EXTRACT] Batch using existing OCR text for {filename}: {len(existing_ocr_text)} chars")
            else:
                # Extract text from binary file before workflow
                if file_data:
                    try:
                        extracted_text = self._extract_text_from_file(file_data, filename)
                        logger.info(f"[EXTRACT] Batch extracted text for {filename}: {len(extracted_text)} chars")
                        
                        # PART 5: FAIL EXPLICITLY ON EMPTY EXTRACTION
                        if not extracted_text.strip():
                            logger.error(f"[EXTRACT] Batch empty text extraction for {filename}")
                            # Create error result for this document
                            error_result = DocumentProcessingResult(
                                request_id=str(uuid.uuid4()),
                                success=False,
                                processing_time_ms=0,
                                processing_type=doc.get('processing_type', processing_type),
                                filename=filename,
                                data=None,
                                metadata={},
                                error_message="Unable to extract text from document - file may be corrupted or unreadable"
                            )
                            processed_documents.append(error_result)
                            continue
                        
                        # PART 2: PASS EXTRACTED TEXT INTO OPTIONS
                        options["ocr_text"] = extracted_text
                        logger.info(f"[EXTRACT] Batch passing extracted text for {filename}")
                        
                    except Exception as extraction_error:
                        logger.error(f"[EXTRACT] Batch text extraction failed for {filename}: {extraction_error}")
                        
                        # CHECK IF THIS IS AN OCR FAILURE (IMAGE FILE)
                        file_extension = filename.lower().split('.')[-1] if '.' in filename.lower() else ''
                        is_image_file = file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']
                        error_text = str(extraction_error)
                        is_ocr_failure = (
                            "OCR failed" in error_text
                            or "Tesseract" in error_text
                            or "OCR_FAILURE" in error_text
                        )
                        
                        if is_image_file and is_ocr_failure:
                            # OCR FAILURE: Return proper response with risk_flags
                            logger.error(f"[OCR] failed for image {filename}: {extraction_error}")
                            
                            # Create OCR failure response data
                            ocr_failure_data = {
                                "document_type": "unknown",
                                "structured_data": {},
                                "extracted_fields": {},
                                "missing_fields": [],
                                "confidence": 0.0,
                                "reasoning": ["OCR failed - image OCR unavailable or extraction failed"],
                                "classification_confidence": 0.0,
                                "classification_reasoning": {
                                    "keywords_found": [],
                                    "structural_indicators": [],
                                    "confidence_factors": ["OCR failed - image OCR unavailable or extraction failed"]
                                },
                                "risk_flags": [
                                    {
                                        "code": "OCR_FAILURE",
                                        "severity": "high",
                                        "message": "OCR failed - image OCR unavailable or extraction failed"
                                    }
                                ],
                                "decision_support": {
                                    "decision": "manual_review_required",
                                    "confidence": 0.0,
                                    "reasoning": ["OCR failed - image OCR unavailable or extraction failed"]
                                },
                                "canonical": {},
                                "summary": "Automatic extraction failed. Manual review required.",
                                "ai_summary": "Automatic extraction failed. Manual review required."
                            }
                            
                            ocr_failure_result = DocumentProcessingResult(
                                request_id=str(uuid.uuid4()),
                                success=True,  # Success=true but with OCR failure flags
                                processing_time_ms=0,
                                processing_type=doc.get('processing_type', processing_type),
                                filename=filename,
                                data=ocr_failure_data,
                                metadata={"ocr_failure": True},
                                error_message=None
                            )
                            processed_documents.append(ocr_failure_result)
                            continue
                        else:
                            # REGULAR EXTRACTION FAILURE: Return normal error
                            error_result = DocumentProcessingResult(
                                request_id=str(uuid.uuid4()),
                                success=False,
                                processing_time_ms=0,
                                processing_type=doc.get('processing_type', processing_type),
                                filename=filename,
                                data=None,
                                metadata={},
                                error_message=f"Text extraction failed: {str(extraction_error)}"
                            )
                            processed_documents.append(error_result)
                            continue
                else:
                    logger.warning(f"[EXTRACT] Batch no file data for {filename}")
                    # Create error result for this document
                    error_result = DocumentProcessingResult(
                        request_id=str(uuid.uuid4()),
                        success=False,
                        processing_time_ms=0,
                        processing_type=doc.get('processing_type', processing_type),
                        filename=filename,
                        data=None,
                        metadata={},
                        error_message="No file data provided for text extraction"
                    )
                    processed_documents.append(error_result)
                    continue
            
            # Add processed document to batch
            processed_documents.append(doc)
        
        # Separate successfully processed documents from errors
        successful_docs = [doc for doc in processed_documents if not isinstance(doc, DocumentProcessingResult)]
        error_results = [doc for doc in processed_documents if isinstance(doc, DocumentProcessingResult)]
        
        # Process successful documents through processor
        batch_results = []
        if successful_docs:
            processor_results = self.processor.process_batch_documents(successful_docs)
            batch_results = [self._convert_to_result_format(result) for result in processor_results]
        
        # Combine error results with processor results
        all_results = error_results + batch_results
        
        return all_results
    
    async def classify_document(
        self,
        file_data: bytes,
        filename: str,
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Classify document type
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            options: Additional options
        
        Returns:
            Document classification result
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Ensure options exists
        options = options or {}
        
        # PART 1: REQUIRED TEXT EXTRACTION ENTRY STEP
        # If valid OCR text already exists, use it; otherwise extract from file
        existing_ocr_text = options.get("ocr_text")
        if existing_ocr_text and isinstance(existing_ocr_text, str) and existing_ocr_text.strip():
            logger.info(f"[EXTRACT] Using existing OCR text for classification: {len(existing_ocr_text)} chars")
        else:
            # Extract text from binary file before classification
            if file_data:
                try:
                    extracted_text = self._extract_text_from_file(file_data, filename)
                    logger.info(f"[EXTRACT] Extracted text for classification: {len(extracted_text)} chars")
                    
                    # PART 5: FAIL EXPLICITLY ON EMPTY EXTRACTION
                    if not extracted_text.strip():
                        logger.error(f"[EXTRACT] Empty text extraction for classification {filename}")
                        return {
                            "document_type": "unknown",
                            "classification_confidence": 0.0,
                            "classification_reasoning": {},
                            "processing_time_ms": 0,
                            "request_id": str(uuid.uuid4()),
                            "error_message": "Unable to extract text from document for classification"
                        }
                    
                    # PART 2: PASS EXTRACTED TEXT INTO OPTIONS
                    options["ocr_text"] = extracted_text
                    logger.info(f"[EXTRACT] Passing extracted text to classification")
                    
                except Exception as extraction_error:
                    logger.error(f"[EXTRACT] Text extraction failed for classification: {extraction_error}")
                    
                    # CHECK IF THIS IS AN OCR FAILURE (IMAGE FILE)
                    file_extension = filename.lower().split('.')[-1] if '.' in filename.lower() else ''
                    is_image_file = file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']
                    error_text = str(extraction_error)
                    is_ocr_failure = (
                        "OCR failed" in error_text
                        or "Tesseract" in error_text
                        or "OCR_FAILURE" in error_text
                    )
                    
                    if is_image_file and is_ocr_failure:
                        # OCR FAILURE: Return proper response with risk_flags
                        logger.error(f"[OCR] failed for image {filename}: {extraction_error}")
                        return {
                            "document_type": "unknown",
                            "classification_confidence": 0.0,
                            "classification_reasoning": {
                                "keywords_found": [],
                                "structural_indicators": [],
                                "confidence_factors": ["OCR failed"]
                            },
                            "processing_time_ms": 0,
                            "request_id": str(uuid.uuid4()),
                            "risk_flags": [
                                {
                                    "code": "OCR_FAILURE",
                                    "severity": "high",
                                    "message": "OCR failed - image OCR unavailable or extraction failed"
                                }
                            ],
                            "decision_support": {
                                "decision": "manual_review_required",
                                "confidence": 0.0,
                                "reasoning": ["OCR failed"]
                            }
                        }
                    else:
                        # REGULAR EXTRACTION FAILURE: Return normal error
                        return {
                            "document_type": "unknown",
                            "classification_confidence": 0.0,
                            "classification_reasoning": {},
                            "processing_time_ms": 0,
                            "request_id": str(uuid.uuid4()),
                            "error_message": f"Text extraction failed: {str(extraction_error)}"
                        }
            else:
                logger.warning("[EXTRACT] No file data provided for classification extraction")
                return {
                    "document_type": "unknown",
                    "classification_confidence": 0.0,
                    "classification_reasoning": {},
                    "processing_time_ms": 0,
                    "request_id": str(uuid.uuid4()),
                    "error_message": "No file data provided for text extraction"
                }
        
        # PART 3: CALL PROCESSOR WITH EXTRACTED TEXT
        result = self.processor.process_document_workflow(
            file_data, filename, "classify", options
        )
        
        if result.get("success", False) and result.get("data"):
            return {
                "document_type": result["data"].get("document_type", "unknown"),
                "classification_confidence": result["data"].get("classification_confidence", 0.0),
                "classification_reasoning": result["data"].get("classification_reasoning", {
                    "keywords_found": [],
                    "structural_indicators": [],
                    "confidence_factors": []
                }),
                "processing_time_ms": result.get("processing_time_ms", 0),
                "request_id": result.get("request_id", "")
            }
        else:
            return {
                "document_type": "unknown",
                "classification_confidence": 0.0,
                "classification_reasoning": {},
                "processing_time_ms": result.get("processing_time_ms", 0),
                "request_id": result.get("request_id", ""),
                "error_message": result.get("error_message", "Classification failed")
            }
    
    def get_supported_processing_types(self) -> List[Dict[str, str]]:
        """Get supported document processing types"""
        return [
            {
                "type": "full_process",
                "description": "Run complete document processing workflow including classification, extraction, validation, risk flags, and decision support"
            },
            {
                "type": "extract_structured",
                "description": "Classify and extract structured fields without full risk/decision analysis"
            },
            {
                "type": "classify",
                "description": "Classify document type only"
            }
        ]
    
    def get_processing_statistics(self, results: List[DocumentProcessingResult]) -> Dict[str, Any]:
        """Get processing statistics for batch results"""
        # Convert DocumentProcessingResult to processor format
        processor_results = []
        for result in results:
            processor_result = {
                'success': result.success,
                'processing_time_ms': result.processing_time_ms,
                'data': {"confidence": result.data.confidence} if result.data else None,
                'error_message': result.error_message
            }
            processor_results.append(processor_result)
        
        return self.processor.get_processing_statistics(processor_results)
    
    def validate_request(self, request: DocumentProcessingRequest) -> Dict[str, Any]:
        """
        Validate document processing request
        
        Args:
            request: Document processing request
        
        Returns:
            Validation result
        """
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # Validate processing type
        supported_types = ["full_process", "extract_structured", "classify"]
        if request.processing_type not in supported_types:
            validation_result["valid"] = False
            validation_result["errors"].append(f"Unsupported processing type: {request.processing_type}")
        
        # Validate options
        if request.options:
            if 'max_file_size' in request.options and request.options['max_file_size'] > 50 * 1024 * 1024:
                validation_result["warnings"].append("Large file size may affect processing performance")
        
        return validation_result
    
    async def process_request(self, request: DocumentProcessingRequest) -> DocumentProcessingResult:
        """
        Process document processing request (main API method)
        
        Args:
            request: Document processing request
        
        Returns:
            DocumentProcessingResult
        """
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info(f"[DOC] metadata request received: {request.processing_type}")
        
        # Extract file URL from options for metadata processing
        options = request.options or {}
        file_url = options.get("fileUrl") or options.get("file_url")
        filename = options.get("filename") or options.get("fileName", "uploaded_document")
        
        # CRITICAL GUARD: Prevent processing with missing fileUrl
        # This prevents duplicate calls where first call succeeds and second call with null fileUrl overwrites success
        if not file_url or not str(file_url).strip():
            logger.error("[DOC] Missing fileUrl in request, returning error")
            print("[DOC] Missing fileUrl - this prevents overwriting successful extraction")
            return DocumentProcessingResult(
                request_id=str(uuid.uuid4()),
                success=False,
                processing_time_ms=0,
                processing_type=request.processing_type,
                filename=filename,
                data=None,
                metadata={},
                error_message="Missing fileUrl for document processing"
            )
        
        # Add required debug log for fileUrl receipt
        print(f"[DOC] received fileUrl: {file_url}")
        logger.info(f"[DOC] fileUrl received: {file_url}")
        
        logger.info(f"[DOC] file fetch started for {filename}")
        
        # Validate request
        validation = self.validate_request(request)
        if not validation["valid"]:
            logger.error(f"[DOC] validation failed: {validation['errors']}")
            return DocumentProcessingResult(
                request_id=str(uuid.uuid4()),
                success=False,
                processing_time_ms=0,
                processing_type=request.processing_type,
                filename=(request.options or {}).get("filename"),
                data=None,
                metadata={},
                error_message="; ".join(validation["errors"])
            )
        
        # Process the document - if we have fileUrl, fetch and process the file
        if file_url:
            try:
                # Import here to avoid circular imports
                import requests
                import os
                from urllib.parse import urlparse
                
                # Download file from URL
                logger.info(f"[DOC] downloading file from {file_url}")
                print("[DOC] downloading file...")
                response = requests.get(file_url, timeout=30)
                response.raise_for_status()
                
                file_content = response.content
                logger.info(f"[DOC] file fetch success, size: {len(file_content)} bytes")
                
                # Extract OCR/text from downloaded file before workflow
                try:
                    extracted_text = self._extract_text_from_file(file_content, filename)
                    logger.info(f"[DOC OCR] extracted text length: {len(extracted_text)}")
                    
                    if not extracted_text.strip():
                        logger.error("[DOC OCR] no text could be extracted from file")
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=False,
                            processing_time_ms=0,
                            processing_type=request.processing_type,
                            filename=filename,
                            data=None,
                            metadata={},
                            error_message="Unable to extract text from downloaded file - file may be corrupted or unreadable"
                        )
                    
                    # Add extracted OCR text to options for workflow
                    options["ocr_text"] = extracted_text
                    logger.info(f"[DOC OCR] passing OCR text to workflow")
                    
                except Exception as ocr_error:
                    logger.error(f"[DOC OCR] extraction failed: {ocr_error}")
                    
                    # CHECK IF THIS IS AN OCR FAILURE (IMAGE FILE)
                    file_extension = filename.lower().split('.')[-1] if '.' in filename.lower() else ''
                    is_image_file = file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']
                    error_text = str(ocr_error)
                    is_ocr_failure = (
                        "OCR failed" in error_text
                        or "Tesseract" in error_text
                        or "OCR_FAILURE" in error_text
                    )
                    
                    if is_image_file and is_ocr_failure:
                        # OCR FAILURE: Return proper response with risk_flags
                        logger.error(f"[OCR] failed for image {filename}: {ocr_error}")
                        
                        # Create OCR failure response data
                        ocr_failure_data = {
                            "document_type": "unknown",
                            "structured_data": {},
                            "extracted_fields": {},
                            "missing_fields": [],
                            "confidence": 0.0,
                            "reasoning": ["OCR failed - image OCR unavailable or extraction failed"],
                            "classification_confidence": 0.0,
                            "classification_reasoning": {
                                "keywords_found": [],
                                "structural_indicators": [],
                                "confidence_factors": ["OCR failed - image OCR unavailable or extraction failed"]
                            },
                            "risk_flags": [
                                {
                                    "code": "OCR_FAILURE",
                                    "severity": "high",
                                    "message": "OCR failed - image OCR unavailable or extraction failed"
                                }
                            ],
                            "decision_support": {
                                "decision": "manual_review_required",
                                "confidence": 0.0,
                                "reasoning": ["OCR failed - image OCR unavailable or extraction failed"]
                            },
                            "canonical": {},
                            "summary": "Automatic extraction failed. Manual review required.",
                            "ai_summary": "Automatic extraction failed. Manual review required."
                        }
                        
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=True,  # Success=true but with OCR failure flags
                            processing_time_ms=0,
                            processing_type=request.processing_type,
                            filename=filename,
                            data=ocr_failure_data,
                            metadata={"ocr_failure": True},
                            error_message=None
                        )
                    else:
                        # REGULAR EXTRACTION FAILURE: Return normal error
                        return DocumentProcessingResult(
                            request_id=str(uuid.uuid4()),
                            success=False,
                            processing_time_ms=0,
                            processing_type=request.processing_type,
                            filename=filename,
                            data=None,
                            metadata={},
                            error_message=f"Text extraction failed: {str(ocr_error)}"
                        )
                
                logger.info(f"[DOC] extraction output: processing with workflow")
                
                # Process the downloaded file content with OCR text
                processing_result = self.processor.process_document_workflow(
                    file_content, filename, request.processing_type, options
                )
                
                # Add required debug log for extraction result
                print(f"[DOC] extraction result: {processing_result}")
                logger.info(f"[DOC] extraction result: {processing_result}")
                
                # Ensure extraction result is not empty and return explicit error if it fails
                if not processing_result or not processing_result.get("data"):
                    logger.error("[DOC] extraction returned empty result")
                    return DocumentProcessingResult(
                        request_id=str(uuid.uuid4()),
                        success=False,
                        processing_time_ms=0,
                        processing_type=request.processing_type,
                        filename=filename,
                        data=None,
                        metadata={},
                        error_message="Document extraction returned empty result - OCR text may be empty or handler failed"
                    )
                
                # Check if OCR text is empty (this is often the root cause)
                extracted_data = processing_result.get("data", {})
                if extracted_data.get("document_type") == "unknown" and not extracted_data.get("structured_data") and not extracted_data.get("extracted_fields"):
                    logger.error("[DOC] extraction failed - OCR text may be empty or unreadable")
                    return DocumentProcessingResult(
                        request_id=str(uuid.uuid4()),
                        success=False,
                        processing_time_ms=0,
                        processing_type=request.processing_type,
                        filename=filename,
                        data=None,
                        metadata={},
                        error_message="Document extraction failed - OCR text may be empty or unreadable"
                    )
                
            except Exception as fetch_error:
                logger.error(f"[DOC] file fetch failed: {fetch_error}")
                # File download failed - return explicit failure
                return DocumentProcessingResult(
                    request_id=str(uuid.uuid4()),
                    success=False,
                    processing_time_ms=0,
                    processing_type=request.processing_type,
                    filename=filename,
                    data=None,
                    metadata={},
                    error_message=f"File download failed: {str(fetch_error)}"
                )
        else:
            # No file URL provided - process with empty content (may work for some workflows)
            logger.warning("[DOC] no file URL provided, processing with empty content")
            file_content = b""
            
            processing_result = self.processor.process_document_workflow(
                file_content, filename, request.processing_type, options
            )
        
        # Convert and return result
        result = self._convert_to_result_format(processing_result)
        
        if result.success:
            logger.info(f"[DOC] extraction success: {result.data is not None}")
            if result.data:
                logger.info(f"[DOC] extraction output keys: {list(result.data.keys()) if isinstance(result.data, dict) else 'non-dict'}")
        else:
            logger.error(f"[DOC] extraction failed: {result.error_message}")
        
        return result
    
    def _convert_to_result_format(self, processing_result: Dict[str, Any]) -> DocumentProcessingResult:
        """
        Convert processor result to DocumentProcessingResult format
        STABILIZED: Single normalized data flow, no mixed dict/object access
        
        Args:
            processing_result: Raw processor result
        
        Returns:
            DocumentProcessingResult model
        """
        logger = logging.getLogger(__name__)
        
        # STEP 1: Normalize input to safe dict at the top
        normalized_result = self._safe_dict(processing_result)
        data = self._safe_dict(normalized_result.get("data"))
        
        # STEP 2: Build extracted_data once with all required fields
        extracted_data = {
            "document_type": data.get("document_type", "unknown"),
            "structured_data": self._safe_dict(data.get("structured_data")),
            "extracted_fields": self._safe_dict(data.get("extracted_fields")),
            "missing_fields": self._safe_list(data.get("missing_fields")),
            "confidence": self._safe_number(data.get("confidence") or data.get("extraction_confidence", 0)),
            "reasoning": self._safe_list(data.get("reasoning")),
            "canonical": self._safe_dict(data.get("canonical")),
            "classification_confidence": self._safe_number(data.get("classification_confidence", 0)),
            "classification_reasoning": self._safe_list(data.get("classification_reasoning"))
        }
        
        # STEP 3: Generate intelligence outputs with safe fallbacks
        intelligence_outputs = self._generate_intelligence_outputs(extracted_data)
        
        # STEP 4: Merge intelligence outputs into both extracted_data and response data
        extracted_data.update(intelligence_outputs)
        data.update(intelligence_outputs)
        
        # STEP 5: Apply ML analysis with deterministic fallbacks
        ml_outputs = self._apply_ml_analysis(extracted_data, data)
        extracted_data.update(ml_outputs)
        data.update(ml_outputs)
        
        # STEP 6: Apply supporting document sanitization if needed
        if extracted_data.get("document_type") == "supporting_document":
            extracted_data = self._sanitize_supporting_document(extracted_data)
            data["structured_data"] = extracted_data.get("structured_data", {})
            data["extracted_fields"] = extracted_data.get("extracted_fields", {})
            
            # Re-apply intelligence outputs for supporting documents
            data.update({k: intelligence_outputs.get(k) for k in ["summary", "ai_summary", "case_insight", "decision_support", "predictions", "decision"] if k in intelligence_outputs})
            
            # Apply AI analysis for supporting documents
            ai_outputs = self._apply_ai_analysis(extracted_data)
            data.update(ai_outputs)
        
        # STEP 7: Final validation and fixes
        self._apply_final_validation(data, extracted_data)
        
        # STEP 7.5: Apply canonical builder and extracted fields normalizer
        # Build canonical from structured_data
        canonical_from_structured = self._build_canonical_data(data)
        existing_canonical = self._safe_dict(data.get("canonical"))
        
        # Merge canonical data - prefer structured_data canonical but preserve valid existing values
        final_canonical = canonical_from_structured.copy()
        for key, value in existing_canonical.items():
            if value is not None and str(value).strip() and key not in final_canonical:
                final_canonical[key] = value
        
        data["canonical"] = final_canonical
        extracted_data["canonical"] = final_canonical
        
        # Normalize extracted_fields
        data["extracted_fields"] = self._normalize_extracted_fields(data)
        extracted_data["extracted_fields"] = data["extracted_fields"]
        
        # STEP 8: Apply deterministic verification routing
        self._authoritative_final_evaluation(extracted_data, data)
        
        # STEP 9: Build final result
        return DocumentProcessingResult(
            request_id=normalized_result.get("request_id", ""),
            success=normalized_result.get("success", False),
            processing_time_ms=normalized_result.get("processing_time_ms", 0),
            processing_type=normalized_result.get("processing_type", ""),
            filename=normalized_result.get("filename", ""),
            data=data,
            metadata=normalized_result.get("metadata", {}),
            error_message=normalized_result.get("error_message")
        )
    def _generate_intelligence_outputs(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate intelligence outputs with safe fallbacks"""
        logger = logging.getLogger(__name__)
        
        # Initialize outputs with safe defaults
        summary = ""
        case_insight = []
        decision_support = {}
        predictions = {}
        
        # Generate each intelligence component independently
        try:
            tmp_summary = self.intelligence_service.generate_document_summary(extracted_data)
            summary = self._safe_str(tmp_summary)
        except Exception as ie:
            logger.warning(f"[INTELLIGENCE] generate_document_summary failed: {ie}")
            
        try:
            tmp_case_insight = self.intelligence_service.generate_case_insight(extracted_data)
            case_insight = self._safe_list(tmp_case_insight)
        except Exception as ie:
            logger.warning(f"[INTELLIGENCE] generate_case_insight failed: {ie}")
            
        try:
            tmp_decision_support = self.intelligence_service.generate_decision_support(extracted_data)
            decision_support = self._safe_dict(tmp_decision_support)
        except Exception as ie:
            logger.warning(f"[INTELLIGENCE] generate_decision_support failed: {ie}")
            
        try:
            tmp_predictions = self.intelligence_service.generate_predictions(extracted_data)
            predictions = self._safe_dict(tmp_predictions)
        except Exception as ie:
            logger.warning(f"[INTELLIGENCE] generate_predictions failed: {ie}")
        
        # Apply deterministic fallbacks
        if not summary:
            summary = self._build_officer_summary(extracted_data)
            
        # Normalize decision_support
        if not decision_support:
            decision_support = {"decision": "review", "confidence": 0.5, "reasoning": []}
        else:
            decision_support.setdefault("decision", "review")
            decision_support.setdefault("confidence", 0.5)
            if not isinstance(decision_support.get("reasoning"), list):
                decision_support["reasoning"] = self._safe_list(decision_support.get("reasoning"))
        
        if not predictions:
            predictions = {}
            
        # Build final intelligence outputs
        final_decision = decision_support.get("decision", "review")
        
        return {
            "summary": summary,
            "ai_summary": summary,  # PART 3: Single source for frontend
            "case_insight": case_insight,
            "decision_support": decision_support,
            "predictions": predictions,
            "decision": final_decision
        }

    def _apply_ml_analysis(self, extracted_data: Dict[str, Any], data: Dict[str, Any]) -> Dict[str, Any]:
        """Apply ML analysis with deterministic fallbacks"""
        logger = logging.getLogger(__name__)
        
        ml_insights = {}
        predictions = {}
        
        # RANDOM FOREST ML INTEGRATION (separate, defensive)
        if ML_SERVICE_AVAILABLE:
            try:
                ml_analysis = None
                try:
                    ml_analysis = analyze_document_risk(extracted_data)
                except Exception as ml_error:
                    logger.warning(f"[ML] Random Forest ML call failed: {ml_error}")
                    ml_analysis = {}

                ml_analysis = self._safe_dict(ml_analysis)

                # Safe nested fields
                pte = self._safe_dict(ml_analysis.get("processing_time_estimate", {}))
                priority_score = self._safe_number(ml_analysis.get("priority_score", 50.0), default=50.0)
                queue = ml_analysis.get("queue") if isinstance(ml_analysis.get("queue"), str) else "NORMAL"
                risk_level = ml_analysis.get("risk_level") if isinstance(ml_analysis.get("risk_level"), str) else "Medium"
                processing_time_days = pte.get("estimated_days") if isinstance(pte.get("estimated_days"), (int, float)) else 2
                approval_likelihood_num = self._safe_number(ml_analysis.get("approval_likelihood", 50), default=50)
                approval_likelihood = f"{int(round(approval_likelihood_num))}%"
                model_confidence = self._safe_number(ml_analysis.get("confidence_score", 0.5), default=0.5)
                auto_decision = ml_analysis.get("auto_decision") if isinstance(ml_analysis.get("auto_decision"), str) else "manual_review"

                ml_insights = {
                    "priority_score": priority_score,
                    "queue": queue,
                    "risk_level": risk_level,
                    "processing_time": processing_time_days,
                    "approval_likelihood": approval_likelihood,
                    "model_confidence": model_confidence,
                    "auto_decision": auto_decision,
                    "prediction_method": "random_forest"
                }

                # If ML signals high risk, add reasoning and escalate if appropriate
                if str(risk_level).lower() == "high":
                    decision_support = data.get("decision_support", {})
                    if not isinstance(decision_support.get("reasoning"), list):
                        decision_support["reasoning"] = self._safe_list(decision_support.get("reasoning"))
                    decision_support["reasoning"].append("High risk detected by ML model")
                    # Escalate to manual review only if not an explicit approve
                    if decision_support.get("decision") != "approve":
                        decision_support["decision"] = "manual_review"
                        extracted_data["decision"] = "manual_review"
                        data["decision"] = "manual_review"
                        data["decision_support"] = decision_support

                predictions = {
                    "priority_score": priority_score,
                    "queue": queue,
                    "risk_level": risk_level,
                    "model_confidence": model_confidence,
                    "prediction_method": "random_forest",
                    "processing_time": f"{processing_time_days} days",
                    "approval_likelihood": approval_likelihood
                }

            except Exception as ml_error:
                logger.warning(f"[ML] Random Forest ML error: {ml_error}")
                # Fall back to deterministic priority computation
                self._fallback_ml_priority(extracted_data, data, data.get("decision_support", {}))
        else:
            # No ML service available -> deterministic fallback
            self._fallback_ml_priority(extracted_data, data, data.get("decision_support", {}))
        
        return {
            "ml_insights": ml_insights,
            "predictions": predictions
        }

    def _apply_ai_analysis(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Apply AI analysis for supporting documents"""
        logger = logging.getLogger(__name__)
        
        # Run LLM analysis defensively
        try:
            ai_output = analyze_application(extracted_data)
        except Exception as ai_err:
            logger.warning(f"[AI] analyze_application failed: {ai_err}")
            ai_output = {}

        ai_output = self._safe_dict(ai_output)

        # Add AI analysis to the main data object
        ai_updates = {
            "ai_summary": self._safe_str(ai_output.get("ai_summary", ""))
        }

        # Merge AI risk flags with existing ones (AI flags come first)
        ai_risk_flags = self._safe_list(ai_output.get("risk_flags", []))
        existing_risk_flags = self._safe_list(ai_updates.get("risk_flags", []))
        ai_updates["risk_flags"] = ai_risk_flags + existing_risk_flags

        # FIX 1: SINGLE SOURCE OF TRUTH - Deterministic decision_support ALWAYS takes precedence
        deterministic_decision_support = self._safe_dict(extracted_data.get("decision_support"))
        ai_decision_support = self._safe_dict(ai_output.get("decision_support"))

        if deterministic_decision_support:
            # Ensure top-level decision matches deterministic decision_support (SINGLE SOURCE OF TRUTH)
            ai_updates["decision"] = deterministic_decision_support.get("decision", "review")
        elif ai_decision_support:
            # Only use AI decision_support if deterministic one is absent
            ai_updates["decision_support"] = ai_decision_support
            # Sync top-level decision with AI decision_support
            ai_updates["decision"] = ai_decision_support.get("decision", "review")
        
        # STEP 7: LLM REFINEMENT LAYER - Add advisory block without overriding deterministic truth
        try:
            extracted_data_with_refinement = add_llm_refinement_to_response(extracted_data)
            
            # Add llm_refinement to main data object
            ai_updates["llm_refinement"] = extracted_data_with_refinement.get("llm_refinement", {})
                
        except Exception as e:
            # LLM refinement failure should not break the pipeline
            fallback_refinement = {
                "refined_summary": "LLM refinement temporarily unavailable",
                "officer_note": "Proceed with deterministic analysis",
                "consistency_flags": [],
                "confidence_note": f"LLM refinement failed: {str(e)}"
            }
            
            ai_updates["llm_refinement"] = fallback_refinement

        return ai_updates

    def _fallback_ml_priority(self, extracted_data: Dict[str, Any], data: Dict[str, Any], decision_support: Dict[str, Any]) -> None:
        """
        Deterministic ML priority fallback when ML service is unavailable
        PART 6: Stabilize ML priority without replacing model
        """
        logger = logging.getLogger(__name__)
        
        try:
            # Normalize input features for deterministic computation
            document_type = extracted_data.get("document_type", "unknown")
            confidence = extracted_data.get("confidence", 0.0)
            missing_fields = extracted_data.get("missing_fields", [])
            structured_data = extracted_data.get("structured_data", {})
            
            # Count identity and amount fields
            identity_fields = ["farmer_name", "applicant_name", "aadhaar_number", "mobile_number"]
            identity_count = sum(1 for field in identity_fields if structured_data.get(field))
            
            amount_fields = ["requested_amount", "claim_amount", "subsidy_amount", "amount"]
            has_amount = any(structured_data.get(field) for field in amount_fields)
            
            # Count risk flags
            risk_flags = extracted_data.get("risk_flags", [])
            risk_count = len(risk_flags) if isinstance(risk_flags, list) else 0
            
            # Deterministic priority score (0-100)
            # Base score from confidence
            priority_score = confidence * 50
            
            # Bonus for identity fields
            priority_score += identity_count * 10
            
            # Bonus for amount presence
            if has_amount:
                priority_score += 15
            
            # Penalty for missing fields
            priority_score -= len(missing_fields) * 5
            
            # Penalty for risk flags
            priority_score -= risk_count * 3
            
            # Clamp to valid range
            priority_score = max(0, min(100, priority_score))
            
            # Deterministic queue assignment
            if confidence < 0.5 or len(missing_fields) > 3 or risk_count > 2:
                queue = "VERIFICATION_QUEUE"
            elif priority_score > 70:
                queue = "LOW"
            elif priority_score > 40:
                queue = "NORMAL"
            else:
                queue = "HIGH"
            
            # Deterministic risk level
            if risk_count > 2 or confidence < 0.3:
                risk_level = "High"
            elif risk_count > 0 or confidence < 0.6:
                risk_level = "Medium"
            else:
                risk_level = "Low"
            
            # Deterministic auto decision
            if confidence < 0.4 or len(missing_fields) > 4 or risk_count > 3:
                auto_decision = "manual_review"
            elif confidence > 0.8 and len(missing_fields) == 0 and risk_count == 0:
                auto_decision = "auto_approve"
            else:
                auto_decision = "review"
            
            # Build fallback ML insights
            fallback_ml = {
                "priority_score": priority_score,
                "queue": queue,
                "risk_level": risk_level,
                "processing_time": 3 if queue == "VERIFICATION_QUEUE" else 2,
                "approval_likelihood": max(20, 80 - risk_count * 10 - len(missing_fields) * 5),
                "model_confidence": confidence,
                "auto_decision": auto_decision,
                "prediction_method": "deterministic_fallback"
            }
            
            # Update both data structures
            data["ml_insights"] = fallback_ml
            data["predictions"] = fallback_ml.copy()
            
            logger.info(f"[ML] Deterministic fallback applied: priority={priority_score}, queue={queue}, risk={risk_level}")
            
        except Exception as e:
            logger.error(f"[ML] Deterministic fallback failed: {e}")
            # Ultra-safe fallback
            safe_fallback = {
                "priority_score": 50.0,
                "queue": "NORMAL",
                "risk_level": "Medium",
                "processing_time": 2,
                "approval_likelihood": 50,
                "model_confidence": 0.5,
                "auto_decision": "review",
                "prediction_method": "safe_fallback"
            }
            data["ml_insights"] = safe_fallback
            data["predictions"] = safe_fallback.copy()

    def _apply_final_validation(self, data: Dict[str, Any], extracted_data: Dict[str, Any]) -> None:
        """Apply final validation and fixes to ensure data consistency"""
        logger = logging.getLogger(__name__)
        
        # STRICT VALIDATION: Ensure all required intelligence fields are present
        if data and isinstance(data, dict):
            # Convert to dict for easier validation
            data_dict = data
            
            # AMOUNT VALIDATION: preserve realistic amounts, remove only clearly invalid ones
            structured_data = data_dict.get("structured_data", {})
            amount_fields = ["requested_amount", "claim_amount", "subsidy_amount", "amount"]
            invalid_amount_removed = False
            for amt_field in amount_fields:
                if amt_field in structured_data and structured_data.get(amt_field) is not None and str(structured_data.get(amt_field)).strip():
                    raw_val = structured_data.get(amt_field)
                    try:
                        # Clean numeric value
                        clean = str(raw_val).replace(",", "").replace("₹", "").replace("$", "").strip()
                        import re
                        clean = re.sub(r"[^0-9.]", "", clean)
                        if clean:
                            amount_num = float(clean)
                            # PART 7: Enhanced amount preservation - only remove clearly unrealistic values
                            # Preserve small valid amounts (75, 8500, 11200)
                            # Only remove: < 1, > 10 million, or obvious junk
                            if amount_num < 1 or amount_num > 10000000:
                                # Additional check: if it's a small integer, it might be valid
                                if amount_num < 1 and raw_val.isdigit() and len(str(raw_val)) <= 5:
                                    # Keep small integers like "75", "850" as they could be valid amounts
                                    logger.info(f"[AMOUNT] Preserving small integer amount: {raw_val}")
                                    continue
                                
                                # Remove invalid amount from structured_data
                                del structured_data[amt_field]
                                data["structured_data"] = structured_data
                                invalid_amount_removed = True
                                logger.info(f"[AMOUNT] Removed invalid amount: {raw_val} (parsed as {amount_num})")
                            else:
                                # Valid amount - ensure proper formatting
                                structured_data[amt_field] = amount_num
                                logger.info(f"[AMOUNT] Preserved valid amount: {raw_val} -> {amount_num}")
                    except Exception as e:
                        # If parsing fails, check if it's a simple integer
                        if str(raw_val).isdigit() and len(str(raw_val)) <= 5:
                            logger.info(f"[AMOUNT] Preserving unparsed integer: {raw_val}")
                        else:
                            logger.warning(f"[AMOUNT] Could not parse amount: {raw_val} - {e}")
                            # Don't remove unparsable amounts that might be valid text

            if invalid_amount_removed:
                # Add explicit reasoning about removed invalid amount(s)
                data.setdefault("reasoning", [])
                data["reasoning"].append("Invalid amount removed")

            # DYNAMIC SCORE: compute score from confidence and missing fields
            try:
                doc_type_for_score = data_dict.get("document_type", "unknown")
                required_fields_list = self.extraction_service._compute_missing_fields(doc_type_for_score, {})
                required_count = len(required_fields_list) if required_fields_list is not None else 0
                missing_count = len(data_dict.get("missing_fields", []))
                confidence_val = data_dict.get("confidence", data_dict.get("extraction_confidence", 0.0)) or 0.0
                missing_ratio = missing_count / max(required_count, 1)
                score_val = confidence_val * (1 - missing_ratio)
                # Clamp score between 0 and 1
                score_val = max(0.0, min(1.0, score_val))
                # Persist score
                data["score"] = score_val
            except Exception:
                pass

            # Validate summary is meaningful (not generic)
            summary = data_dict.get("summary")
            if not summary or "temporarily unavailable" in summary or "scheme application regarding scheme application" in summary:
                # Build a minimal, non-hallucinated officer-facing summary using only extracted fields
                document_type = data_dict.get("document_type", "document")
                structured_data = data_dict.get("structured_data", {}) or {}
                name = structured_data.get("farmer_name") or structured_data.get("applicant_name") or structured_data.get("applicant") or "Unknown"

                fixed_summary = f"Applicant {name} submitted a {document_type.replace('_', ' ')} related request."

                # If an amount exists (and wasn't removed), include it
                amt_val = None
                for af in amount_fields:
                    if af in structured_data and structured_data.get(af) is not None and str(structured_data.get(af)).strip():
                        amt_val = structured_data.get(af)
                        break
                if amt_val:
                    # Format numeric amounts cleanly when possible
                    try:
                        clean_amt = str(amt_val).replace(",", "").replace("₹", "").replace("$", "").strip()
                        import re
                        clean_amt_digits = re.sub(r"[^0-9.]", "", clean_amt)
                        if clean_amt_digits:
                            amt_num = float(clean_amt_digits)
                            if amt_num >= 1:
                                if amt_num >= 100000:
                                    formatted = f"₹{int(amt_num):,}"
                                else:
                                    formatted = f"₹{int(amt_num)}"
                                fixed_summary += f" The requested amount is {formatted}."
                            else:
                                fixed_summary += f" The requested amount is {amt_val}."
                        else:
                            fixed_summary += f" The requested amount is {amt_val}."
                    except Exception:
                        fixed_summary += f" The requested amount is {amt_val}."

                # Include a short list of key fields if present
                important_keys = ["scheme_name", "application_id", "aadhaar_number", "village", "location"]
                key_parts = []
                for k in important_keys:
                    if k in structured_data and structured_data.get(k):
                        key_parts.append(f"{k}: {structured_data.get(k)}")
                if key_parts:
                    fixed_summary += f" Key details include {', '.join(key_parts)}."

                data["summary"] = fixed_summary
                data["ai_summary"] = fixed_summary
            
            # ENSURE SUMMARY SYNC: Always keep ai_summary in sync with summary
            if data.get("summary"):
                data["ai_summary"] = data["summary"]
            
            # Validate decision exists
            decision = data_dict.get("decision")
            if not decision or decision not in ["approve", "review", "reject", "manual_review"]:
                # Auto-fix with conservative decision
                data["decision"] = "review"
            
            # Validate ml_insights exists
            ml_insights = data_dict.get("ml_insights")
            if not ml_insights or not isinstance(ml_insights, dict):
                logger.warning("[ML] ml_insights validation failed, using fallback")
                # Auto-fix with fallback ML insights using real deterministic method
                ml_prediction = predict_application_priority(data_dict)
                fallback_ml = {
                    "priority_score": ml_prediction.get("priority_score", 0.4),
                    "queue": ml_prediction.get("queue", "LOW"),
                    "review_type": ml_prediction.get("review_type", "MANUAL_REVIEW"),
                    "model_confidence": ml_prediction.get("model_confidence", 0.4),
                    "prediction_method": ml_prediction.get("prediction_method", "rule_based_fallback")
                }
                # Add generic predictions for non-ML fields
                fallback_ml.update({
                    "processing_time": "3 days",
                    "approval_likelihood": "50%",
                    "risk_level": "Medium"
                })
                data["ml_insights"] = fallback_ml
                data["predictions"] = fallback_ml  # Also update predictions!
            else:
                logger.info("[ML] ml_insights validation passed, keeping existing")
                # Ensure predictions field matches ml_insights prediction_method
                data["predictions"] = {
                    "priority_score": ml_insights.get("priority_score", 0.5),
                    "queue": ml_insights.get("queue", "NORMAL"),
                    "review_type": ml_insights.get("review_type", "AUTO"),
                    "model_confidence": ml_insights.get("model_confidence", 0.5),
                    "prediction_method": ml_insights.get("prediction_method", "trained_random_forest")
                }
                # Add generic predictions for non-ML fields
                data["predictions"].update({
                    "processing_time": ml_insights.get("processing_time", "2 days"),
                    "approval_likelihood": ml_insights.get("approval_likelihood", "50%"),
                    "risk_level": ml_insights.get("risk_level", "Medium")
                })
            
            # Validate case_insight has at least 3 useful lines
            case_insight = data_dict.get("case_insight", [])
            if not case_insight or len(case_insight) < 3:
                # Auto-fix with basic case insight
                structured_data = data_dict.get("structured_data", {})
                farmer_name = structured_data.get("farmer_name", "Unknown farmer") if structured_data else "Unknown farmer"
                document_type = data_dict.get("document_type", "document")
                
                fixed_insight = [
                    f"Farmer: {farmer_name}",
                    f"Request: {document_type.replace('_', ' ')}",
                    "Status: Requires review"
                ]
                
                data["case_insight"] = fixed_insight
            
            # WORKFLOW INTEGRATION: Add workflow information
            try:
                workflow_summary = self.workflow_service.get_workflow_summary(data_dict)
                data["workflow"] = workflow_summary["workflow"]
                data["next_steps"] = workflow_summary["next_steps"]
                data["sla_timeline"] = workflow_summary["sla_timeline"]
            except Exception as e:
                logger.warning(f"[WORKFLOW] Workflow integration failed: {e}")
                # Add fallback workflow
                fallback_workflow = {
                    "status": "PENDING_REVIEW",
                    "queue": "NORMAL",
                    "estimated_processing_time": "2 days",
                    "requires_manual_review": True,
                    "risk_level": "Medium"
                }
                data["workflow"] = fallback_workflow
                data["next_steps"] = ["Standard processing", "Review required"]
                data["sla_timeline"] = {"initial_review": "3-5 business days", "final_decision": "7-10 business days"}

    def _authoritative_final_evaluation(self, extracted_data: Dict[str, Any], data: Any):
        """Single authoritative evaluation to ensure missing_fields, risk_flags, and reasoning are consistent"""
        
        document_type = extracted_data.get("document_type", "unknown")
        structured_data = extracted_data.get("structured_data", {})
        
        # Define required fields per document type
        required_fields_map = {
            'scheme_application': ['farmer_name', 'scheme_name'],  # REMOVED application_id - not reliably present
            'farmer_record': ['farmer_name', 'aadhaar_number', 'mobile_number'],
            'grievance': ['farmer_name', 'grievance_type', 'description'],
            'insurance_claim': ['farmer_name', 'claim_type', 'loss_description'],
            'subsidy_claim': ['farmer_name', 'subsidy_type', 'requested_amount'],
            'supporting_document': [],
            'unknown': []
        }
        
        # Get required fields for this document type
        required_fields = required_fields_map.get(document_type, [])
        
        # Compute missing fields
        missing_fields = []
        for field in required_fields:
            if not structured_data.get(field):
                missing_fields.append(field)
        
        # Update missing fields in extracted data
        extracted_data["missing_fields"] = missing_fields
        
        # Compute confidence based on missing fields
        confidence = 1.0
        if required_fields:
            present_fields = len([f for f in required_fields if structured_data.get(f)])
            confidence = present_fields / len(required_fields)
        
        extracted_data["confidence"] = confidence
        
        # PART 5: Authoritative verification routing based on actual conditions
        decision_support = extracted_data.get("decision_support", {})
        ml_insights = extracted_data.get("ml_insights", {})
        
        # Determine if verification is required based on comprehensive conditions
        requires_verification = (
            confidence < 0.7 or  # Low confidence
            len(missing_fields) > 2 or  # Many missing fields
            decision_support.get("decision") in ["manual_review", "review"] or  # Manual review required
            ml_insights.get("auto_decision") == "manual_review" or  # ML says manual review
            ml_insights.get("risk_level") == "high" or  # High risk
            extracted_data.get("document_type") == "supporting_document" or  # Supporting documents need review
            len(extracted_data.get("risk_flags", [])) > 2 or  # Many risk flags
            confidence < 0.5  # Very low confidence threshold
        )
        
        # Set queue and workflow based on verification requirement
        if requires_verification:
            queue = "VERIFICATION_QUEUE"
            workflow_status = "UNDER_REVIEW"
        else:
            # Use ML queue suggestion but ensure safe defaults
            ml_queue = ml_insights.get("queue", "NORMAL")
            if ml_queue in ["HIGH", "NORMAL", "LOW"]:
                queue = ml_queue
            else:
                queue = "NORMAL"  # Safe default
            workflow_status = "CASE_READY"
        
        # Update workflow information
        workflow = extracted_data.get("workflow", {})
        workflow["queue"] = queue
        workflow["status"] = workflow_status
        workflow["requires_manual_review"] = requires_verification
        
        extracted_data["workflow"] = workflow
        
        # Add routing reasoning to extracted data for transparency
        routing_reasons = []
        if confidence < 0.7:
            routing_reasons.append(f"Low confidence ({confidence:.2f})")
        if len(missing_fields) > 2:
            routing_reasons.append(f"Many missing fields ({len(missing_fields)})")
        if decision_support.get("decision") in ["manual_review", "review"]:
            routing_reasons.append(f"Decision support requires {decision_support.get('decision')}")
        if ml_insights.get("auto_decision") == "manual_review":
            routing_reasons.append("ML recommends manual review")
        if ml_insights.get("risk_level") == "high":
            routing_reasons.append("High risk level detected")
        if extracted_data.get("document_type") == "supporting_document":
            routing_reasons.append("Supporting document requires review")
        
        if routing_reasons:
            extracted_data["routing_reasoning"] = routing_reasons

    # --- Text extraction methods ----------------------------------------------

    def _sanitize_supporting_document(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Sanitize supporting document data to remove financial fields and contaminated metadata
        
        Args:
            extracted_data: Raw extracted data from supporting document
            
        Returns:
            Sanitized extracted data
        """
        sanitized = extracted_data.copy()
        
        # Remove financial fields that may contain identifier-like numbers
        financial_fields_to_remove = [
            "requested_amount", "claim_amount", "subsidy_amount", 
            "loan_amount", "amount", "subsidy_amount"
        ]
        
        structured_data = sanitized.get("structured_data", {})
        extracted_fields = sanitized.get("extracted_fields", {})
        
        # Remove from structured_data
        for field in financial_fields_to_remove:
            structured_data.pop(field, None)
            extracted_fields.pop(field, None)
        
        # Remove contaminated person_name fields that are generic placeholders
        generic_name_patterns = [
            "supporting document", "document", "proof", "receipt", 
            "applicant", "beneficiary", "claimant", "holder"
        ]
        
        # Check and remove generic names
        for field_name in ["farmer_name", "applicant_name", "claimant_name", "complainant"]:
            if field_name in structured_data:
                value = str(structured_data[field_name]).lower().strip()
                if any(pattern in value for pattern in generic_name_patterns):
                    structured_data.pop(field_name, None)
                    extracted_fields.pop(field_name, None)
            
            if field_name in extracted_fields:
                value = str(extracted_fields[field_name]).lower().strip()
                if any(pattern in value for pattern in generic_name_patterns):
                    extracted_fields.pop(field_name, None)
        
        # Remove contaminated location/address fields with metadata markers
        contamination_patterns = [
            "supporting document metadata", "field value", "supporting document type",
            "reference number", "document type", "document meta"
        ]
        
        location_fields = ["location", "village", "district", "address", "place"]
        for field_name in location_fields:
            if field_name in structured_data:
                value = str(structured_data[field_name]).lower()
                if any(pattern in value for pattern in contamination_patterns):
                    structured_data.pop(field_name, None)
                    extracted_fields.pop(field_name, None)
                    
            if field_name in extracted_fields:
                value = str(extracted_fields[field_name]).lower()
                if any(pattern in value for pattern in contamination_patterns):
                    extracted_fields.pop(field_name, None)
        
        # Remove contaminated document_type_detail with multiline junk
        if "document_type_detail" in structured_data:
            value = str(structured_data["document_type_detail"])
            if "\\n" in value or any(pattern in value.lower() for pattern in contamination_patterns):
                structured_data.pop("document_type_detail", None)
                extracted_fields.pop("document_type_detail", None)
                
        if "document_type_detail" in extracted_fields:
            value = str(extracted_fields["document_type_detail"])
            if "\\n" in value or any(pattern in value.lower() for pattern in contamination_patterns):
                extracted_fields.pop("document_type_detail", None)
        
        # Update sanitized data
        sanitized["structured_data"] = structured_data
        sanitized["extracted_fields"] = extracted_fields
        
        return sanitized
    
    def _extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        Unified text extraction from file based on type
        Supports: PDF, Images (JPG/PNG), DOCX, TXT
        
        Args:
            file_content: Raw file bytes
            filename: Original filename for type detection
            
        Returns:
            Clean extracted text (never empty)
            
        Raises:
            ValueError: If extraction fails or returns empty text
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Detect file type from filename
        file_extension = filename.lower().split('.')[-1] if '.' in filename.lower() else ''
        logger.info(f"[EXTRACT] File type detected: {file_extension}")
        
        # Handle .doc files explicitly with error
        if file_extension == 'doc':
            raise ValueError("Legacy .doc files are not supported. Please upload PDF, DOCX, or image.")
        
        try:
            text = ""
            
            if file_extension == 'pdf':
                text = self._extract_text_from_pdf(file_content)
            elif file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
                text = self._extract_text_from_image(file_content, filename)
            elif file_extension == 'docx':
                text = self._extract_text_from_docx(file_content)
            elif file_extension in ['txt', 'text']:
                text = self._extract_text_from_text_file(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Normalize text
            text = (text or "").strip()
            
            # Strong validation - fail if empty
            if not text:
                raise ValueError("Extraction returned empty text")
            
            # For images, check for insufficient text
            if file_extension in ['jpg', 'jpeg', 'png', 'tiff', 'bmp'] and len(text) < 10:
                raise ValueError("Extraction returned insufficient text")
            
            return text
            
        except Exception as e:
            logger.error(f"[EXTRACT] Extraction error for {file_extension}: {e}")
            # Always raise - no silent fallback
            raise ValueError(f"Text extraction failed: {str(e)}")
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF bytes using pdfplumber"""
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            import pdfplumber  # type: ignore[reportMissingImports]
            import io
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                            logger.debug(f"[DOC OCR] extracted {len(page_text)} chars from page {page_num + 1}")
                    except Exception as page_error:
                        logger.warning(f"[DOC OCR] failed to extract page {page_num + 1}: {page_error}")
                        continue
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"[DOC OCR] PDF extraction complete: {len(extracted_text)} chars from {len(text_parts)} pages")
            return extracted_text
            
        except Exception as e:
            logger.error(f"[DOC OCR] PDF extraction failed: {e}")
            # Fallback to PyMuPDF if pdfplumber fails
            try:
                import fitz  # PyMuPDF  # type: ignore[reportMissingImports]
                doc = fitz.open(stream=file_content, filetype="pdf")
                text_parts = []
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                
                doc.close()
                extracted_text = "\n".join(text_parts)
                logger.info(f"[DOC OCR] PyMuPDF fallback extraction: {len(extracted_text)} chars")
                return extracted_text
                
            except Exception as fallback_error:
                logger.error(f"[DOC OCR] PyMuPDF fallback also failed: {fallback_error}")
                raise
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes using python-docx"""
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            import docx  # type: ignore[reportMissingImports]
            import io
            
            # Create a temporary file-like object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Open the document
            doc = docx.Document(docx_file)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"[EXTRACT] DOCX extraction complete: {len(extracted_text)} chars")
            return extracted_text
            
        except ImportError:
            logger.error("[EXTRACT] python-docx not installed, cannot extract from DOCX")
            raise ValueError("DOCX extraction not available - python-docx library not installed")
        except Exception as e:
            logger.error(f"[EXTRACT] DOCX extraction failed: {e}")
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_text_from_image(self, file_content: bytes, filename: str) -> str:
        """Extract text from image bytes using Tesseract OCR."""
        import logging
        logger = logging.getLogger(__name__)

        import io
        from PIL import Image
        import PIL.ImageEnhance
        import PIL.ImageFilter
        import pytesseract

        tesseract_path = self._resolve_tesseract_binary()
        if not tesseract_path:
            raise ValueError("OCR_FAILURE: Tesseract binary not found")

        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        logger.info(f"[OCR] Using Tesseract at: {tesseract_path}")
        
        try:
            image = PIL.Image.open(io.BytesIO(file_content))

            if image.mode != 'RGB':
                image = image.convert('RGB')

            enhancer = PIL.ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)

            gray_image = image.convert('L')

            try:
                import numpy as np
                img_array = np.array(gray_image)
                threshold = 128
                img_array = np.where(img_array > threshold, 255, 0).astype(np.uint8)
                processed_image = Image.fromarray(img_array, mode='L')
            except ImportError:
                processed_image = gray_image.point(lambda x: 0 if x < 128 else 255, mode='1')

            try:
                tesseract_config = get_tesseract_config()
                language_candidates = self._get_ocr_language_candidates()

                extracted_text = ""
                selected_language = ""
                last_language_error = None

                for language_code in language_candidates:
                    try:
                        if tesseract_config:
                            extracted_text = pytesseract.image_to_string(
                                processed_image,
                                lang=language_code,
                                config=tesseract_config,
                            )
                        else:
                            extracted_text = pytesseract.image_to_string(
                                processed_image,
                                lang=language_code,
                            )
                        selected_language = language_code
                        break
                    except Exception as lang_error:
                        last_language_error = lang_error
                        logger.warning(f"[OCR] OCR failed with {language_code}, trying fallback: {lang_error}")

                if not extracted_text.strip() and last_language_error is not None:
                    raise last_language_error

                logger.info(f"[OCR] OCR completed for {filename} using {selected_language or 'eng'}")
            except Exception as lang_error:
                logger.warning(f"[OCR] OCR language fallback failed for {filename}: {lang_error}")
                if get_tesseract_config():
                    extracted_text = pytesseract.image_to_string(processed_image, lang="eng", config=get_tesseract_config())
                else:
                    extracted_text = pytesseract.image_to_string(processed_image, lang="eng")

            text_length = len(extracted_text.strip())

            if text_length == 0:
                raise ValueError("OCR failed: no text could be extracted from image")

            if text_length < 20:
                logger.warning(f"[OCR] Very low text extraction from image {filename}: {text_length} chars")
                extracted_text = f"[LOW_OCR_QUALITY] {extracted_text.strip()}"

            ocr_error_patterns = [
                "|||||||", "||||||", "|||||", "||||",  # Vertical lines (common in handwritten forms)
                "_____", "____", "___",  # Underlines
                "[illegible]", "[unreadable]", "[unclear]"  # Explicit illegibility markers
            ]

            for pattern in ocr_error_patterns:
                if pattern in extracted_text.lower():
                    logger.warning(f"[OCR] Detected handwriting pattern '{pattern}' in {filename}")
                    extracted_text = f"[HANDWRITING_DETECTED] {extracted_text.strip()}"
                    break

            logger.info(f"[OCR] Image OCR extraction complete: {text_length} chars")
            return extracted_text.strip()

        except Exception as ocr_error:
            logger.error(f"[OCR] Tesseract OCR processing failed: {ocr_error}")
            raise ValueError(f"OCR_FAILURE: Tesseract extraction failed - {str(ocr_error)}")
    
    def _extract_text_from_text_file(self, file_content: bytes) -> str:
        """Extract text from text file bytes"""
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"[DOC OCR] Text file extraction complete with {encoding}: {len(text)} chars")
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            text = file_content.decode('utf-8', errors='ignore')
            logger.warning(f"[DOC OCR] Text file extraction with errors: {len(text)} chars")
            return text.strip()
            
        except Exception as e:
            logger.error(f"[DOC OCR] Text file extraction failed: {e}")
            raise
